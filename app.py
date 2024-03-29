import streamlit as st
from openai import OpenAI
from docx import Document
import os
import tempfile

client = OpenAI(
    # defaults to os.environ.get("OPENAI_API_KEY")
    api_key = "My_API_Key",
)

def transcribe_audio(audio_file):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as temp_audio:
        temp_audio_path = temp_audio.name
        temp_audio.write(audio_file.read())
    
    try:
        with open(temp_audio_path, 'rb') as audio_file:
            transcription = client.audio.transcriptions.create(model="whisper-1", file=audio_file)
        return transcription.text
    finally:
        os.remove(temp_audio_path)

def meeting_minutes(transcription):
    abstract_summary = abstract_summary_extraction(transcription)
    key_points = key_points_extraction(transcription)
    action_items = action_item_extraction(transcription)
    sentiment = sentiment_analysis(transcription)
    return {
        'abstract_summary': abstract_summary,
        'key_points': key_points,
        'action_items': action_items,
        'sentiment': sentiment
    }


def abstract_summary_extraction(transcription):
    response = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a highly skilled AI trained in language comprehension and summarization. I would like you to read the following text and summarize it into a concise abstract paragraph. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content



def key_points_extraction(transcription):
    response = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a proficient AI with a specialty in distilling information into key points. Based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content


def action_item_extraction(transcription):
    response = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are an AI expert in analyzing conversations and extracting action items. Please review the text and identify any tasks, assignments, or actions that were agreed upon or mentioned as needing to be done. These could be tasks assigned to specific individuals, or general actions that the group has decided to take. Please list these action items clearly and concisely."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content


def sentiment_analysis(transcription):
    response = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "As an AI with expertise in language and emotion analysis, your task is to analyze the sentiment of the following text. Please consider the overall tone of the discussion, the emotion conveyed by the language used, and the context in which words and phrases are used. Indicate whether the sentiment is generally positive, negative, or neutral, and provide brief explanations for your analysis where possible."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content


# Function to save minutes as a docx file
def save_as_docx(minutes, filename):
    doc = Document()
    for key, value in minutes.items():
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        doc.add_paragraph()  # Add a line break between sections
    doc.save(filename)

# Streamlit app
st.title("Meeting Minutes Generator")

# File upload widget
uploaded_file = st.file_uploader("Upload an audio file (mp3)", type=["mp3"])

if uploaded_file is not None:
    # Display uploaded file details
    st.subheader("Uploaded Audio File:")
    st.write(f"File Name: {uploaded_file.name}")
    st.write(f"File Type: {uploaded_file.type}")

    # Transcribe audio and generate minutes
    with st.spinner("Transcribing audio..."):
        transcription = transcribe_audio(uploaded_file)

    minutes = meeting_minutes(transcription)

    # Display minutes
    st.subheader("Meeting Minutes:")
    for key, value in minutes.items():
        st.write(f"**{key.replace('_', ' ').title()}**: {value}")

    # Save minutes as a docx file
    st.subheader("Download Minutes as Docx:")
    download_button = st.button("Download Docx")
    if download_button:
        with st.spinner("Saving as Docx..."):
            save_as_docx(minutes, "meeting_minutes.docx")
        st.success("Docx file downloaded successfully!")
